"""
DOCX / PDF extractor.

Capabilities:
- Extract raw text from DOCX files (python-docx)
- Extract raw text from PDF files (pdfplumber or PyMuPDF)
- Locate likely sections for regulatory fields using heuristic pattern matching
- Return structured candidate extractions with source snippets
- Does NOT fabricate absent values

The returned ExtractionResult carries raw_text_snippets keyed by section type.
Downstream mappers use these snippets for LLM-assisted extraction.
"""

from __future__ import annotations

import logging
import re
from typing import Any, Dict, List, Optional

from ..contracts import CanonicalType, ExtractionResult

logger = logging.getLogger(__name__)


# Heuristic section patterns keyed by logical section name
_SECTION_PATTERNS: Dict[str, List[str]] = {
    "device_description": [
        r"device\s+description",
        r"product\s+description",
        r"general\s+description",
        r"description\s+of\s+(the\s+)?device",
    ],
    "intended_purpose": [
        r"intended\s+purpose",
        r"intended\s+use",
        r"purpose\s+of\s+(the\s+)?device",
        r"medical\s+purpose",
    ],
    "indications_for_use": [
        r"indication[s]?\s+for\s+use",
        r"clinical\s+indication[s]?",
        r"intended\s+indication[s]?",
    ],
    "contraindications": [
        r"contraindication[s]?",
        r"contra-indication[s]?",
        r"warnings\s+and\s+precaution[s]?",
    ],
    "notified_body": [
        r"notified\s+body",
        r"nb\s+(name|number|id)",
        r"conformity\s+assessment\s+body",
    ],
    "eu_mdr_classification": [
        r"(mdr\s+)?classification\s+(rule|class)",
        r"device\s+class",
        r"class\s+(i+[ab]?|iia|iib|iii)",
    ],
    "udi": [
        r"(basic\s+)?udi[-\s]?di",
        r"universal\s+device\s+identifier",
        r"udi\s+number",
    ],
    "pmcf_details": [
        r"post[-\s]?market\s+clinical\s+follow[-\s]?up",
        r"pmcf\s+(plan|report|activities|study)",
        r"clinical\s+follow[-\s]?up",
    ],
    "literature_review": [
        r"literature\s+(review|search|surveillance)",
        r"systematic\s+(review|search)",
        r"published\s+(literature|data|evidence)",
    ],
    "rmf_reference": [
        r"risk\s+management\s+(file|plan|report)",
        r"iso\s+14971",
        r"rmf\s+(number|document|reference)",
    ],
    "ifu_reference": [
        r"instructions?\s+for\s+use",
        r"ifu\s+(number|document|reference)",
        r"user\s+manual",
    ],
    "certificate_details": [
        r"(ce\s+)?certificate\s+(number|date|ref)",
        r"declaration\s+of\s+conformity",
        r"doc\s+(number|date|reference)",
        r"certification\s+date",
    ],
    "sterility": [
        r"steril(e|ity|ization|isation)",
        r"single[- ]use",
        r"reusable",
    ],
}

_SNIPPET_CHARS = 400  # chars to capture around a section match


def _extract_text_docx(path: str) -> str:
    try:
        from docx import Document  # type: ignore
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
        return "\n".join(parts)
    except ImportError:
        logger.warning("python-docx not installed. pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX text extraction failed for {path}: {e}")
        return ""


def _extract_text_pdf(path: str) -> str:
    # Try pdfplumber first, then PyMuPDF
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(path) as pdf:
            return "\n".join(
                page.extract_text() or "" for page in pdf.pages
            )
    except ImportError:
        pass
    except Exception as e:
        logger.error(f"pdfplumber failed for {path}: {e}")

    try:
        import fitz  # type: ignore  # PyMuPDF
        doc = fitz.open(path)
        text = "\n".join(page.get_text() for page in doc)  # type: ignore
        doc.close()
        return text
    except ImportError:
        logger.warning(
            "Neither pdfplumber nor PyMuPDF installed. PDF extraction unavailable. "
            "pip install pdfplumber  OR  pip install pymupdf"
        )
        return ""
    except Exception as e:
        logger.error(f"PyMuPDF failed for {path}: {e}")
        return ""


def _locate_sections(text: str) -> Dict[str, str]:
    """Find section snippets by heuristic pattern matching."""
    snippets: Dict[str, str] = {}
    text_lower = text.lower()
    for section_name, patterns in _SECTION_PATTERNS.items():
        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                start = max(0, match.start() - 30)
                end = min(len(text), match.end() + _SNIPPET_CHARS)
                snippets[section_name] = text[start:end].strip()
                break  # first matching pattern wins
    return snippets


def read_document(
    path: str,
    canonical_type: CanonicalType = CanonicalType.UNKNOWN,
) -> ExtractionResult:
    """
    Extract text and locate sections from a DOCX or PDF file.
    raw_text_snippets carries {section_name -> snippet} for downstream LLM extraction.
    rows[0] carries {"full_text": <first 8000 chars>} for general use.
    """
    ext = path.rsplit(".", 1)[-1].lower()
    notes: List[str] = []

    if ext in ("docx", "doc"):
        text = _extract_text_docx(path)
    elif ext == "pdf":
        text = _extract_text_pdf(path)
    elif ext == "txt":
        try:
            with open(path, encoding="utf-8") as fh:
                text = fh.read()
        except Exception as e:
            text = ""
            notes.append(f"TXT read error: {e}")
    else:
        text = ""
        notes.append(f"Unsupported document format: {ext}")

    if not text:
        notes.append("No text could be extracted from document.")

    snippets = _locate_sections(text) if text else {}
    if snippets:
        notes.append(f"Located sections: {', '.join(snippets.keys())}")
    else:
        notes.append("No known sections located by heuristic matching.")

    rows: List[Dict[str, Any]] = [{"full_text": text[:8000]}] if text else []

    return ExtractionResult(
        source_path=path,
        source_type=canonical_type,
        headers=list(snippets.keys()),
        rows=rows,
        header_row_index=None,
        sheet_name=None,
        extraction_notes=notes,
        raw_text_snippets=snippets,
    )
